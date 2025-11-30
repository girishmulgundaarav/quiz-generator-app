from docx import Document

def create_docx(quiz_data, topic, difficulty):
    doc = Document()
    doc.add_heading(f"Quiz: {topic} ({difficulty})", level=1)

    for i, q in enumerate(quiz_data, start=1):
        doc.add_paragraph(f"{i}. {q['question']}")
        for option in q["options"]:
            doc.add_paragraph(f"  • {option}")

    file_path = "quiz.docx"
    doc.save(file_path)

    return file_path
